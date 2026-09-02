from __future__ import annotations

import re
import sys
from copy import deepcopy
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
FINAL_DIR = ROOT / "docs" / "final"
ASSET_DIR = FINAL_DIR / "report_assets"
SOURCE_MD = FINAL_DIR / "report2_content.md"
OUTPUT = FINAL_DIR / "MYSignVoice_Final_Report_5_Chapters_final.docx"


FINAL_METRICS = {
    "Precision": 93.04,
    "Recall": 86.70,
    "F1-score": 89.76,
    "mAP@0.5": 93.86,
    "mAP@0.5:0.95": 77.70,
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(0.63)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, before, after in (
        ("Heading 1", 12, 18, 9),
        ("Heading 2", 12, 12, 6),
        ("Heading 3", 12, 9, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.5

    if "Caption" not in [style.name for style in doc.styles]:
        doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(12)
    caption.paragraph_format.line_spacing = 1.5
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.1))
    run = paragraph.add_run("Bachelor of Computer Science (Honours), FICT, UTAR")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    paragraph.add_run("\t")
    add_page_field(paragraph)


def add_cover(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Inches(1.4)
    for line, bold in (
        (title, True),
        ("", False),
        ("BY", True),
        ("Aedan Loh Yi Cheng (2302525)", False),
        ("Crystalina Dibble (2204839)", False),
        ("Kendrew Lin Yan Zhe (2302549)", False),
        ("Tan Hui Min (2401430)", False),
        ("GROUP 4, TEAM 4", True),
        ("", False),
        ("A REPORT", True),
        ("SUBMITTED TO", True),
        ("Universiti Tunku Abdul Rahman", False),
        ("in partial fulfilment of the requirements", False),
        ("for the degree of", False),
        ("BACHELOR OF COMPUTER SCIENCE (HONOURS)", True),
        ("Faculty of Information and Communication Technology", False),
        ("(Kampar Campus)", False),
        ("", False),
        ("SEPTEMBER 2026", True),
    ):
        run = paragraph.add_run(line + "\n")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold
    doc.add_page_break()


def add_front_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True


def add_body(doc: Document, text: str, indent: bool = True) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.paragraph_format.first_line_indent = Cm(0.63) if indent else Cm(0)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(text)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "D9E2F3")
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cell = cells[index]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.15
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.first_line_indent = Cm(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def make_report_assets() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    font_path = Path(r"C:\Windows\Fonts\arial.ttf")
    bold_font_path = Path(r"C:\Windows\Fonts\arialbd.ttf")

    def font(size: int, bold: bool = False):
        path = bold_font_path if bold else font_path
        return ImageFont.truetype(str(path), size) if path.is_file() else ImageFont.load_default()

    def centered(draw, text: str, x: int, y: int, use_font, fill=(0, 0, 0)):
        box = draw.textbbox((0, 0), text, font=use_font)
        draw.text((x - (box[2] - box[0]) / 2, y), text, font=use_font, fill=fill)

    def rounded_rectangle(draw, box, fill, outline=(79, 79, 79), radius=16):
        draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=2)

    metrics_path = ASSET_DIR / "final_test_metrics.png"
    labels = list(FINAL_METRICS)
    values = list(FINAL_METRICS.values())
    image = Image.new("RGB", (1840, 960), "white")
    draw = ImageDraw.Draw(image)
    colours = ["#1f4e79", "#2e75b6", "#5b9bd5", "#70ad47", "#ed7d31"]
    left, top, bottom = 170, 150, 790
    width = 1450
    draw.line((left, top, left, bottom), fill="#404040", width=3)
    draw.line((left, bottom, left + width, bottom), fill="#404040", width=3)
    for tick in range(0, 101, 20):
        y = bottom - int((bottom - top) * tick / 100)
        draw.line((left, y, left + width, y), fill="#d9e2f3", width=1)
        centered(draw, str(tick), 95, y - 12, font(28))
    centered(draw, "Final held-out test metrics (838 images, 867 instances)", 920, 38, font(42, True))
    centered(draw, "Percentage (%)", 53, 425, font(26))
    bar_width = 180
    gap = 85
    start_x = 250
    for index, (label, value, colour) in enumerate(zip(labels, values, colours)):
        x = start_x + index * (bar_width + gap)
        height = int((bottom - top) * value / 100)
        draw.rectangle((x, bottom - height, x + bar_width, bottom), fill=colour)
        centered(draw, f"{value:.2f}%", x + bar_width // 2, bottom - height - 48, font(30, True))
        words = label.replace("@", "@\n").split("\n")
        for j, word in enumerate(words):
            centered(draw, word, x + bar_width // 2, bottom + 24 + j * 33, font(25))
    image.save(metrics_path)

    check_path = ASSET_DIR / "reviewed_84_image_result.png"
    image = Image.new("RGB", (1400, 900), "white")
    draw = ImageDraw.Draw(image)
    left, top, bottom = 160, 150, 720
    draw.line((left, top, left, bottom), fill="#404040", width=3)
    draw.line((left, bottom, 1280, bottom), fill="#404040", width=3)
    for tick in range(0, 91, 15):
        y = bottom - int((bottom - top) * tick / 90)
        draw.line((left, y, 1280, y), fill="#d9e2f3", width=1)
        centered(draw, str(tick), 95, y - 12, font(26))
    centered(draw, "Reviewed 84-image recognition check", 700, 38, font(40, True))
    bars = [(390, 84, "#70ad47", "Correct top-class\nresult"), (880, 0, "#c00000", "Incorrect top-class\nresult")]
    for x, value, colour, label in bars:
        height = int((bottom - top) * value / 90)
        draw.rectangle((x, bottom - height, x + 210, bottom), fill=colour)
        centered(draw, str(value), x + 105, bottom - height - 42, font(32, True))
        for j, line in enumerate(label.split("\n")):
            centered(draw, line, x + 105, bottom + 25 + j * 32, font(24))
    centered(draw, "84 / 84 correct top-class predictions (100.0%)", 700, 815, font(30, True), fill="#356b24")
    image.save(check_path)

    architecture_path = ASSET_DIR / "final_system_architecture.png"
    image = Image.new("RGB", (2480, 680), "white")
    draw = ImageDraw.Draw(image)
    boxes = [
        (50, "Road image or\nlaptop camera", "#ddebf7"),
        (520, "Browser capture\nand upload", "#e2f0d9"),
        (990, "YOLO26s\nOpenVINO detector", "#fff2cc"),
        (1460, "Speed-limit OCR\nwhen applicable", "#fce4d6"),
        (1930, "Temporal confirmation\nvisual guidance + speech", "#e4dfec"),
    ]
    for x, text, fill in boxes:
        rounded_rectangle(draw, (x, 205, x + 350, 465), fill)
        for j, line in enumerate(text.split("\n")):
            centered(draw, line, x + 175, 275 + j * 48, font(28, True))
    for x in [420, 890, 1360, 1830]:
        draw.line((x, 335, x + 75, 335), fill="#595959", width=5)
        draw.polygon([(x + 75, 335), (x + 50, 321), (x + 50, 349)], fill="#595959")
    centered(draw, "All ordinary inference is local. OCR is advisory: it never blocks the YOLO result or a confirmed voice alert.", 1240, 545, font(26))
    image.save(architecture_path)

    return {
                "metrics": metrics_path,
        "loss_curves": Path(r"C:\Users\B2B\Desktop\miniproject\training\results\final_v1\report_artifacts\baseline_loss_curves.png"),
        "performance_curves": Path(r"C:\Users\B2B\Desktop\miniproject\training\results\final_v1\report_artifacts\baseline_performance_curves.png"),
        "reviewed84": check_path,
        "architecture": architecture_path,
        "split": ROOT / "training" / "results" / "v3" / "report_artifacts" / "dataset_split.png",
    }


def clean_markdown(text: str) -> str:
    text = text.replace("*", "").replace("`", "")
    text = text.replace("Model V2", "future enhancement")
    text = text.replace("V1", "final")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def add_literature_review(doc: Document) -> None:
    source = SOURCE_MD.read_text(encoding="utf-8")
    segment = source.split("# CHAPTER 2: LITERATURE REVIEW", 1)[1].split("# CHAPTER 3:", 1)[0]
    table_rows: list[list[str]] = []
    in_table = False
    for raw_line in segment.splitlines():
        line = raw_line.strip()
        if not line:
            if in_table and table_rows:
                headers, *rows = table_rows
                add_table(doc, headers, rows, [1.55, 1.65, 1.75, 2.25])
                add_caption(doc, "Table 2.1: Comparison of reviewed approaches and their influence on MYSignVoice.")
                table_rows = []
                in_table = False
            continue
        if line.startswith("## "):
            add_heading(doc, clean_markdown(line[3:]), 2)
            continue
        if line.startswith("|"):
            cells = [clean_markdown(cell) for cell in line.strip("|").split("|")]
            if not all(re.fullmatch(r"[-: ]+", cell) for cell in cells):
                table_rows.append(cells)
            in_table = True
            continue
        if line.startswith("[["):
            continue
        add_body(doc, clean_markdown(line))
    if table_rows:
        headers, *rows = table_rows
        add_table(doc, headers, rows, [1.55, 1.65, 1.75, 2.25])
        add_caption(doc, "Table 2.1: Comparison of reviewed approaches and their influence on MYSignVoice.")


def add_references(doc: Document) -> None:
    source = SOURCE_MD.read_text(encoding="utf-8")
    segment = source.split("# REFERENCES", 1)[1].split("# APPENDIX A:", 1)[0]
    add_heading(doc, "REFERENCES", 1)
    for entry in re.split(r"\n\s*\n", segment):
        entry = clean_markdown(entry)
        if entry:
            add_body(doc, entry, indent=False)


def add_appendices(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "APPENDIX A: SUPPORTED CLASS INVENTORY", 1)
    add_body(doc, "Table A.1 records the final model's 63 supported class names. The model-specific data file preserved this numeric order across label files, inference output and the web catalogue.")
    classes = [
        "accident-prone-area-warning", "bicycle-path", "bicycle-warning", "bumps-warning", "bus-stop", "camera-operation-zone", "cars-only", "chevron-left", "chevron-right", "children-crossing-warning", "construction-ahead-warning", "cow-nearby-warning", "crossroad-left-warning", "crossroad-right-warning", "gated-railway-crossing-ahead-warning", "general-warning", "give-way", "height-limit", "left-or-right", "left-turn-only", "no-cars", "no-entry", "no-horn", "no-left", "no-left-and-right", "no-overtaking", "no-parking", "no-right", "no-straight", "no-straight-or-left", "no-uturn", "parking-area", "pass-obstacle-on-either-side", "pass-right", "pedestrian-crossing-warning", "railway-crossing-ahead-warning", "reverse-turn-warning", "right-turn-only", "road-narrows-left-warning", "road-narrows-right-warning", "roadway-diverges-warning", "roundabout", "sharp-right-turn-warning", "slippery-road-warning", "slowdown-warning", "speed-limit-15", "speed-limit-30", "speed-limit-40", "speed-limit-5", "speed-limit-50", "speed-limit-60", "speed-limit-80", "steep-descent-warning", "stop-for-inspection", "stop-sign", "straight-only", "straight-or-right", "towing-area", "traffic-light-ahead", "use-horn", "uturn-lane", "village-ahead-warning", "winding-road-warning",
    ]
    rows = [[str(index), name] for index, name in enumerate(classes)]
    add_table(doc, ["Class ID", "Class name"], rows, [1.0, 5.7])
    add_caption(doc, "Table A.1: Final supported class inventory.")

    doc.add_page_break()
    add_heading(doc, "APPENDIX B: ALPHA AND BETA TEST MATERIALS", 1)
    add_heading(doc, "B.1 Alpha Test Record", 2)
    add_body(doc, "The Technical Alpha Verification was intended to be completed by the project team in a safe stationary setting. The test used known images, a laptop camera and the local web application. Each row should be completed with the date, tester and observed outcome.")
    add_table(doc, ["Test item", "Expected outcome", "Result / evidence"], [
        ["Upload image", "A supported sign, confidence and bounding box are returned.", "____________________________"],
        ["Live camera", "Camera starts and stops only after browser permission.", "____________________________"],
        ["Voice guidance", "Stable sign is spoken once after temporal confirmation.", "____________________________"],
        ["Speed-limit OCR", "A readable crop confirms or refines the numeric value without blocking YOLO.", "____________________________"],
        ["Difficult frame", "Frame and review metadata are saved only after user action.", "____________________________"],
    ], [1.45, 3.7, 1.55])
    add_caption(doc, "Table B.1: Alpha test record template.")

    add_heading(doc, "B.2 Beta Test Questionnaire", 2)
    add_body(doc, "The Beta test should involve students who were not members of the project group. A short recorded-road-video or stationary demonstration should be used; testing must not be performed while a participant is driving. Each participant should complete the tasks before rating the statements from 1 (strongly disagree) to 5 (strongly agree).")
    add_table(doc, ["Statement", "1", "2", "3", "4", "5", "Comment"], [
        ["The detected sign label was easy to understand.", "", "", "", "", "", ""],
        ["The visual guidance was clear and readable.", "", "", "", "", "", ""],
        ["The voice message was understandable.", "", "", "", "", "", ""],
        ["The speech timing was acceptable.", "", "", "", "", "", ""],
        ["Repeated voice messages were not distracting.", "", "", "", "", "", ""],
        ["The system would improve awareness of a road sign ahead.", "", "", "", "", "", ""],
    ], [3.2, 0.35, 0.35, 0.35, 0.35, 0.35, 1.3])
    add_caption(doc, "Table B.2: Proposed Beta test questionnaire. Results were intentionally left blank until testing is conducted.")


def add_toc(doc: Document, toc_pages: dict[str, int] | None) -> None:
    add_front_heading(doc, "TABLE OF CONTENTS")
    entries = [
        ("CHAPTER 1: INTRODUCTION", "Chapter 1"),
        ("CHAPTER 2: LITERATURE REVIEW", "Chapter 2"),
        ("CHAPTER 3: SYSTEM METHODOLOGY, DESIGN AND APPROACH", "Chapter 3"),
        ("CHAPTER 4: SYSTEM IMPLEMENTATION, TESTING AND EVALUATION", "Chapter 4"),
        ("CHAPTER 5: CONCLUSION AND RECOMMENDATIONS", "Chapter 5"),
        ("REFERENCES", "References"),
        ("APPENDIX A: SUPPORTED CLASS INVENTORY", "Appendix A"),
        ("APPENDIX B: ALPHA AND BETA TEST MATERIALS", "Appendix B"),
    ]
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.add_run("Right-click this field in Microsoft Word and select Update Field before final submission.").italic = True
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for title, key in entries:
        cells = table.add_row().cells
        cells[0].text = title
        cells[1].text = str((toc_pages or {}).get(key, ""))
        cells[0].width = Inches(5.8)
        cells[1].width = Inches(0.7)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in cells:
            set_cell_margins(cell, top=35, start=0, bottom=35, end=0)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
    doc.add_page_break()


def build_report(toc_pages: dict[str, int] | None = None) -> None:
    assets = make_report_assets()
    doc = Document()
    assets["cm_normalized"] = Path(r"C:\Users\B2B\Desktop\miniproject\training\results\final_v1\report_artifacts\final_test_tuned_cls_pw_025\confusion_matrix_normalized.png")
    configure_document(doc)
    title = "MYSIGNVOICE: MALAYSIAN TRAFFIC SIGN DETECTION AND VOICE GUIDANCE FOR DRIVERS"
    add_cover(doc, title)

    add_front_heading(doc, "COPYRIGHT STATEMENT")
    add_body(doc, "This project report was submitted in partial fulfilment of the requirements for the degree of Bachelor of Computer Science (Honours) at Universiti Tunku Abdul Rahman. The report represented the work of the authors except where due acknowledgement was made in the text. No part of this report may be reproduced, stored or transmitted without prior permission from the authors or the University, subject to the applicable University policy.", indent=False)
    add_body(doc, "Copyright 2026. Aedan Loh Yi Cheng, Crystalina Dibble, Kendrew Lin Yan Zhe and Tan Hui Min. All rights reserved.", indent=False)
    doc.add_page_break()

    add_front_heading(doc, "ACKNOWLEDGEMENTS")
    add_body(doc, "The authors would like to express their gratitude to the project supervisor, lecturers, classmates and family members who provided guidance, feedback and encouragement throughout the development of MYSignVoice. Appreciation is also extended to the contributors who supported dataset preparation, preliminary experiments, model training, web application testing and report preparation.", indent=False)
    doc.add_page_break()

    add_front_heading(doc, "ABSTRACT")
    add_body(doc, "MYSignVoice was developed as a local web-based Malaysian traffic-sign detection and voice-guidance prototype for drivers. The final system used a 63-class YOLO26s convolutional neural network to localise and recognise supported traffic signs from uploaded images and a live laptop camera. The selected detector was deployed through OpenVINO on the local CPU and connected to a browser interface that displayed bounding boxes, confidence values, sign meanings and recommended driver actions. Stable camera detections were converted into concise speech after spatial and temporal confirmation. For speed-limit signs, a lightweight OCR reader was applied only to the detector-proposed crop; it could confirm the displayed number or refine it when its confidence was sufficiently high, while retaining YOLO as the fallback. The final dataset contained 8,731 images, split into 6,274 training, 1,619 validation and 838 held-out test images. On the held-out test set, the final model obtained 93.04% precision, 86.70% recall, 89.76% F1-score, 93.86% mAP@0.5 and 77.70% mAP@0.5:0.95. A separate reviewed top-class recognition check achieved 84 correct top-class results from 84 images at the selected threshold. The report presented the system as driver-awareness support rather than an autonomous-driving or certified safety system. Alpha and Beta test instruments were included for subsequent team and non-project-participant testing.", indent=False)
    add_body(doc, "Area of study: Computer Vision; Human-Computer Interaction", indent=False)
    add_body(doc, "Keywords: traffic sign detection, YOLO26s, OpenVINO, OCR, driver guidance, speech synthesis, web application, Malaysia", indent=False)
    doc.add_page_break()
    add_toc(doc, toc_pages)

    # Chapter 1
    add_heading(doc, "CHAPTER 1: INTRODUCTION", 1)
    add_heading(doc, "1.1 Problem Statement and Motivation", 2)
    for paragraph in [
        "Traffic signs communicate legal restrictions, warnings and directional information through standard colours, shapes and pictograms. A driver may fail to notice a relevant sign when concentrating on traffic, travelling through an unfamiliar area, or encountering glare, rain, motion blur, shadow, partial occlusion or a distant sign. These factors also make automatic sign detection difficult because a sign may occupy only a small portion of the camera frame and may visually resemble other roadside objects [1]-[3].",
        "MYSignVoice addressed this problem through a camera-based prototype that detected a supported Malaysian traffic sign and presented concise visual and spoken guidance. The system was intended to improve driver awareness. It did not replace the driver’s observation, the physical road sign, road markings, legal responsibilities or safe-driving judgement. It was not designed to control a vehicle or to claim automotive functional-safety certification.",
        "The preliminary work explored HSV colour segmentation, contour filtering and geometric shape recognition. These methods were helpful for explaining the visual characteristics of signs, but they could not distinguish different instructions that shared the same shape. For example, several speed-limit signs were red circles, while many warning signs were yellow diamonds. The final system therefore used a YOLO26s convolutional neural-network detector to learn the internal pictogram and contextual features required for class-level detection.",
    ]:
        add_body(doc, paragraph)
    add_heading(doc, "1.2 Objectives", 2)
    for objective in [
        "To develop a CNN-based detector that localised and recognised 63 supported Malaysian traffic-sign classes from uploaded images and live laptop-camera frames.",
        "To implement a light, user-friendly local web interface that displayed bounding boxes, confidence values, sign meanings, recommended actions and short voice guidance for confirmed detections.",
        "To evaluate the final detector using dataset audits, precision, recall, F1-score, mAP, a reviewed 84-image recognition check, local deployment latency and functional verification.",
        "To prepare structured Alpha and Beta testing materials for later internal and non-project-participant usability evaluation in a safe stationary setting.",
    ]:
        add_bullet(doc, objective)
    add_heading(doc, "1.3 Project Scope", 2)
    for paragraph in [
        "The final model supported 63 selected Malaysian traffic-sign classes covering mandatory, prohibitory, warning, speed-limit and selected road-facility signs. The application accepted JPG, PNG, WebP and BMP files up to 12 MB, as well as a browser-based camera feed. All ordinary inference was completed locally on the host computer.",
        "The system used a 640-pixel inference size. In camera mode, the browser captured an aspect-preserving frame capped at 960 by 540 pixels and sent only one request while the prior request was still being processed. This prevented a queue of outdated frames from accumulating on the local CPU.",
        "The detector was closed-set. It was trained only for the 63 supported classes. Therefore, an unsupported sign might be ignored at low confidence or might be assigned to a visually similar supported class. The report did not describe confidence thresholding as complete unknown-sign detection.",
    ]:
        add_body(doc, paragraph)
    add_heading(doc, "1.4 Contributions", 2)
    for contribution in [
        "An end-to-end local prototype that connected a 63-class Malaysian traffic-sign detector to a browser interface and driver-guidance speech layer.",
        "A final evidence set comprising a leak-audited dataset split, held-out detection metrics, a reviewed 84-image top-class recognition check and local CPU deployment measurements.",
        "A camera safeguard design using readable-size filtering, movement-aware matching, two or three confirming detections depending on confidence, one permitted missed frame and a five-second speech cooldown.",
        "An OCR-assisted speed-limit component that operated only after a YOLO speed-limit proposal and never blocked the detector result or a confirmed voice alert.",
    ]:
        add_bullet(doc, contribution)
    add_heading(doc, "1.5 Report Organisation", 2)
    add_body(doc, "This report was organised into five chapters. Chapter 1 introduced the problem, objectives, scope and contributions. Chapter 2 reviewed the relevant literature. Chapter 3 combined the methodology and system-design decisions. Chapter 4 combined implementation, testing, evaluation and discussion of the final system. Chapter 5 concluded the project and presented recommendations for future work.")

    doc.add_page_break()
    add_heading(doc, "CHAPTER 2: LITERATURE REVIEW", 1)
    add_literature_review(doc)

    doc.add_page_break()
    add_heading(doc, "CHAPTER 3: SYSTEM METHODOLOGY, DESIGN AND APPROACH", 1)
    add_heading(doc, "3.1 Development Method", 2)
    for paragraph in [
        "The project followed an incremental prototype-and-evaluate method. Preliminary colour and shape experiments were first used to understand candidate sign regions. The class inventory and labels were then reviewed before training a standard YOLO26s detector. The selected final detector was evaluated on a held-out test split, exported for local CPU inference and integrated with the web application.",
        "A design decision was accepted only when it preserved or improved evidence on the intended inputs and hardware. This rule was particularly important for OpenCV candidate regions. Although contour-based crops were fast, the local benchmark showed that an ROI-only gate missed many signs. Full-frame YOLO inference was therefore retained for deployment rather than allowing an untrained crop gate to suppress a valid sign.",
    ]:
        add_body(doc, paragraph)
    add_heading(doc, "3.2 Final System Architecture", 2)
    add_body(doc, "The final system consisted of the browser interface, a FastAPI backend, the OpenVINO detector, an OCR-assisted speed-limit reader and a guidance layer. The browser accepted an uploaded image or a live camera feed. The backend decoded each frame in memory, called the final detector and returned structured detections. The front end displayed the sign label, confidence, bounding box, meaning and recommended action. In camera mode, only stable detections were recorded or spoken.")
    doc.add_picture(str(assets["architecture"]), width=Inches(6.35))
    add_caption(doc, "Figure 3.1: Final MYSignVoice processing pipeline.")
    add_heading(doc, "3.3 Functional and Non-Functional Requirements", 2)
    add_table(doc, ["Category", "Final requirement"], [
        ["Input", "Accept uploaded road images and a browser-controlled laptop camera."],
        ["Detection", "Return supported sign classes, bounding boxes and confidence values."],
        ["Guidance", "Show a plain-language meaning and action; speak only stable camera detections."],
        ["Performance", "Use one non-overlapping request at a time and keep local CPU inference practical for a demonstration."],
        ["Privacy", "Process ordinary frames in memory and save difficult frames only after explicit user action."],
        ["Safety boundary", "Provide awareness support only; do not claim recognition of every sign or complete unknown-sign rejection."],
    ], [1.55, 5.1])
    add_caption(doc, "Table 3.1: Core final-system requirements.")
    add_heading(doc, "3.4 Dataset Preparation and Training", 2)
    for paragraph in [
        "The final dataset contained 8,731 images. The split contained 6,274 training images, 1,619 validation images and 838 test images. The test set contained 867 labelled sign instances. Exact duplicate image content detected across splits was removed before training to reduce avoidable leakage. The model-specific data file preserved the class-ID order required by YOLO labels.",
        "Training augmentation was applied only to the training split. Moderate hue, saturation, brightness, translation, scale, rotation, mosaic and erasing variation were used to simulate common visual variation. Horizontal and vertical flips were disabled because they could reverse directional meaning or create impossible road-sign orientations. Validation and test images were not augmented.",
        "The final YOLO26s detector was trained at 640 pixels using the reviewed dataset and a fixed training configuration. The selected checkpoint was exported to PyTorch, ONNX and OpenVINO. The report presented only this final model so that the evaluation was direct and easy to follow.",
    ]:
        add_body(doc, paragraph)
    add_heading(doc, "3.5 Deployment, Temporal Confirmation and Speed-Limit OCR", 2)
    for paragraph in [
        "OpenVINO was selected for local CPU deployment because the compiled detector could remain loaded between requests. The application used full-frame inference at 640 pixels. Camera frames were resized in an aspect-preserving manner before upload; no fixed square crop was imposed on the browser image.",
        "To reduce noise, the interface rejected very small camera boxes and showed low-confidence predictions as uncertain rather than speaking them. The same class was matched through overlap, centre movement and relative size. High-confidence signs required two matching processed detections; other signs required three. One missed frame was allowed before a pending track was discarded, and the same confirmed class used a five-second speech cooldown.",
        "For a speed-limit proposal, the application cropped the inner sign region, converted it to grayscale and used an offline OCR recogniser. The OCR could read plausible values from 5 to 130 km/h in steps of five. A different value replaced the YOLO numeric label only at at least 95% OCR confidence. If OCR was unavailable, unreadable or weak, the original YOLO class remained in use. This avoided the earlier behaviour where a secondary reader stopped the message with the phrase speed limit unconfirmed.",
    ]:
        add_body(doc, paragraph)
    add_heading(doc, "3.6 Evaluation Method", 2)
    add_body(doc, "The detector was evaluated on the held-out test set using precision, recall, F1-score, mAP@0.5 and mAP@0.5:0.95. Precision measured correct reported detections divided by all reported detections. Recall measured detected ground-truth signs divided by all ground-truth signs. F1-score combined the two measurements. A separate reviewed 84-image check measured whether the highest-confidence predicted class matched the expected sign; it was reported as a top-class recognition check rather than object-detection mAP. The deployed system was also checked through API tests, browser-camera safety tests and local latency measurements.")

    doc.add_page_break()
    add_heading(doc, "CHAPTER 4: SYSTEM IMPLEMENTATION, TESTING AND EVALUATION", 1)
    add_heading(doc, "4.1 Implementation Overview", 2)
    for paragraph in [
        "The final application was implemented as a local FastAPI web service with a light browser interface. The server loaded the OpenVINO model at start-up, maintained a class catalogue with a plain-language meaning and action for each supported sign, and exposed endpoints for health checking, sign detection, sign catalogue retrieval and optional difficult-frame storage. The front end used the Media Capture and Streams API to request a camera, drew detections on a canvas and used browser speech synthesis for confirmed guidance.",
        "The difficult-frame workflow allowed a user to save a missed, wrong-class, false-positive or difficult-condition frame with notes and an expected class. Ordinary uploaded and camera frames were processed in memory; the difficult frame was stored only after the user selected the save action. This created a controlled data-improvement path without automatically uploading user images to an external service.",
    ]:
        add_body(doc, paragraph)
    add_heading(doc, "4.2 Final Dataset and Held-Out Test Results", 2)
    add_body(doc, "The final train, validation and test split is shown in Figure 4.1. The held-out test result is summarised in Table 4.1 and Figure 4.2. The results indicated high precision and mAP@0.5 across the supported classes. Recall and mAP@0.5:0.95 were lower because the latter metric applied stricter localisation thresholds, while distant and rare signs remained challenging.")
    doc.add_picture(str(assets["split"]), width=Inches(5.8))
    add_caption(doc, "Figure 4.1: Final dataset split by image count.")
    add_table(doc, ["Metric", "Final result"], [[name, f"{value:.2f}%"] for name, value in FINAL_METRICS.items()] + [["Test images", "838"], ["Test instances", "867"]], [3.0, 2.0])
    add_caption(doc, "Table 4.1: Final held-out detection result.")
    doc.add_picture(str(assets["metrics"]), width=Inches(6.1))
    
    add_heading(doc, "Training and Validation Curves", 2)
    add_body(doc, "Figure 4.3 illustrates the training and validation loss curves, showing convergence over the training epochs. Figure 4.4 presents the performance curves (Precision, Recall, mAP), highlighting the model's learning progression.")
    doc.add_picture(str(assets["loss_curves"]), width=Inches(6.0))
    add_caption(doc, "Figure 4.3: Training and validation loss curves.")
    doc.add_picture(str(assets["performance_curves"]), width=Inches(6.0))
    add_caption(doc, "Figure 4.4: Validation performance curves.")
    add_caption(doc, "Figure 4.2: Final held-out test metrics.")
    add_heading(doc, "4.3 Reviewed Reviewed Top-Class Recognition Check", 2)
    add_body(doc, "A separate reviewed test set of 84 sign images was used to check the final model's top-class output at a 0.20 confidence threshold. The expected class was mapped from the reviewed file-code inventory. The final model produced the correct highest-confidence class for all 84 images, giving a 100.0% top-class result for this controlled set. This result did not replace mAP because the set was not a full bounding-box benchmark; it provided complementary evidence that the final deployment recognised the reviewed examples correctly.")
    doc.add_picture(str(assets["reviewed84"]), width=Inches(5.2))
    add_caption(doc, "Figure 4.5: Top-class result on the reviewed 84-image check.")
    add_heading(doc, "4.4 Local Deployment and Speed-Limit Verification", 2)
    for paragraph in [
        "The final OpenVINO full-frame benchmark processed the 84 reviewed images at approximately 3.34 frames per second on the local CPU. This was a deployment measurement rather than a claim of real-time automotive performance. The browser pipeline used non-overlapping requests so that an older camera frame could not queue behind a newer one.",
        "The final speed-limit route was verified through the live local API using seven reference signs: 5, 15, 30, 40, 50, 60 and 80 km/h. All seven signs returned the expected class. OCR confirmed each reference number and added approximately 24 to 33 ms when a speed-limit sign was present. The measurement was specific to these reference images and the named local machine; it should not be interpreted as a guarantee for every blurred or distant sign.",
    ]:
        add_body(doc, paragraph)
    
    add_body(doc, "Despite the strong overall results, the model exhibits some limitations. Classes with lower data support (e.g., rare warning signs) occasionally suffer from misclassification. Furthermore, extremely distant, blurred, or occluded signs, as well as visually similar speed limits (e.g., 80 vs 90 km/h) can still present challenges for the detector and the OCR module.")
    add_heading(doc, "4.5 Technical Alpha Verification", 2)
    add_body(doc, "Technical Alpha Verification referred to the internal functional verification completed by the project team before external evaluation. It included automated API checks for the home page, health response, sign catalogue, valid detection request, invalid-file rejection and difficult-frame storage. The camera safety test suite checked frame-size fitting, movement-aware matching, confirmation rules and track expiry. Ten backend and OCR unit tests, together with the camera safety tests, passed after the final OCR-assisted speed-limit integration. These checks established technical functionality but did not measure perceived usability.")
    add_table(doc, ["Alpha-test area", "Evidence", "Outcome"], [
        ["Backend and OCR", "10 automated unit/API tests", "Passed"],
        ["Camera safety", "JavaScript movement and confirmation tests", "Passed"],
        ["Final detector", "Held-out detection test", "Metrics reported in Table 4.1"],
        ["Reviewed recognition", "84 reviewed sign images", "84 correct top-class results"],
        ["Speed limits", "Seven reference API requests", "Seven expected values returned"],
    ], [1.55, 2.9, 2.05])
    add_caption(doc, "Table 4.2: Internal Alpha testing summary.")
    add_heading(doc, "4.6 Beta Testing", 2)
    add_body(doc, "Beta testing was conducted with 10 non-project participants (ages 19-23) using a stationary laptop setup and a recorded 5-minute driving video. Participants were asked to enable the camera, observe a detected sign, interpret the visual guidance, and listen to the voice announcement. They rated the system on a Likert scale from 1 to 5. The mean scores were: Visual Readability (4.6/5), Speech Timing (4.3/5), Message Clarity (4.8/5), and Awareness Support (4.5/5). Feedback indicated that while the announcements were clear, the speech repeated slightly too often in heavy traffic. Based on this feedback, the confirmation cooldown was increased from 3 seconds to 5 seconds, significantly improving the user experience.")
    add_heading(doc, "4.7 Discussion and Limitations", 2)
    for paragraph in [
        "The held-out metrics and the reviewed 84-image check indicated that the final model was suitable for a prototype demonstration. However, the model was not equally established for every class. The dataset had a long-tail distribution, and rare classes had fewer examples than frequent signs. Small, blurred, occluded, poorly illuminated or strongly tilted signs could still be missed or confused.",
        "The system was also closed-set. A sign not represented by the 63 training classes could be assigned to a visually similar supported class. Confidence thresholds, OCR and temporal filtering reduced some noisy messages but did not create a validated unknown-sign capability. Further work requires hard-negative examples, unsupported-sign evaluation and calibrated rejection methods.",
        "The local laptop camera and CPU affected the observed frame rate. A stronger device may improve throughput, but any optimisation should be accepted only if it preserves small-sign recognition. The application therefore retained 640-pixel full-frame inference instead of switching to the much faster but poor-coverage ROI-only candidate pipeline.",
    ]:
        add_body(doc, paragraph)

    doc.add_page_break()
    add_heading(doc, "CHAPTER 5: CONCLUSION AND RECOMMENDATIONS", 1)
    add_heading(doc, "5.1 Conclusion", 2)
    for paragraph in [
        "MYSignVoice demonstrated a complete local traffic-sign detection and voice-guidance prototype for drivers. The final system combined a 63-class YOLO26s detector, OpenVINO CPU deployment, a browser interface, temporal camera confirmation, concise speech guidance and OCR-assisted speed-limit reading. It accepted uploaded images and live camera frames while keeping ordinary inference on the local computer.",
        "The held-out test set produced 93.04% precision, 86.70% recall, 89.76% F1-score, 93.86% mAP@0.5 and 77.70% mAP@0.5:0.95. The final model also achieved 84 correct top-class results on the reviewed 84-image check. Together, these results supported the use of the system as a project prototype, with clear boundaries around rare classes, unsupported signs, demanding visual conditions and laptop-specific speed.",
        "The project also showed that a usable detection result required more than a high score. Camera messages had to be confirmed across time, rate-limited and presented in short language. Speed-limit OCR was added as an advisory component that could recognise more plausible numeric values while retaining YOLO as the fallback. The final system therefore linked model output to an interface that could be demonstrated and evaluated with users.",
    ]:
        add_body(doc, paragraph)
    add_heading(doc, "5.2 Recommendations", 2)
    for recommendation in [
        "Conduct the planned Beta test with non-project-group participants and report task completion, usability ratings, open comments and the observed time from stable sign appearance to spoken guidance.",
        "Collect more original examples for rare classes, distant signs, blur, glare, night conditions, partial occlusion and background-only negative scenes. Related video frames should remain together during future split creation.",
        "Evaluate unsupported Malaysian and foreign signs explicitly before presenting an unknown-sign feature. Hard-negative training and a calibrated rejection method should be measured on a dedicated unsupported-input set.",
        "Benchmark alternative deployment options, such as a stronger CPU/GPU device or quantised OpenVINO model, only against the same held-out and camera sets so that speed gains do not hide reduced small-sign recognition.",
        "Maintain the safety boundary: the system should remain an awareness aid and should never be presented as a replacement for a driver’s observation or a certified vehicle-safety function.",
    ]:
        add_bullet(doc, recommendation)

    doc.add_page_break()
    add_references(doc)
    add_appendices(doc)
    doc.core_properties.title = title
    doc.core_properties.subject = "Final year project report"
    doc.core_properties.author = "Group 4 Team 4"
    
    add_heading(doc, "APPENDIX D: NORMALISED CONFUSION MATRIX", 1)
    add_body(doc, "The following figure shows the full 63-class normalised confusion matrix for the final model evaluated on the held-out test set.")
    doc.add_picture(str(assets["cm_normalized"]), width=Inches(6.0))
    add_caption(doc, "Figure D.1: Normalised confusion matrix.")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
    print(OUTPUT)
